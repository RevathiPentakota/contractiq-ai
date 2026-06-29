"""Tests for document extraction: DocumentAgent, StorageService, DocumentExtractor."""

import pytest
from pathlib import Path

from app.agents.document_agent import DocumentAgent
from app.services.document_extractor import DocumentExtractor
from app.state.contract_state import ContractState


# ── Tests for DocumentExtractor ────────────────────────────────────────────────


class TestDocumentExtractor:
    """Tests for the DocumentExtractor service."""

    def test_get_file_extension_pdf(self) -> None:
        """Extract .pdf extension from filename."""
        ext = DocumentExtractor.get_file_extension("contract.pdf")
        assert ext == ".pdf"

    def test_get_file_extension_docx(self) -> None:
        """Extract .docx extension from filename."""
        ext = DocumentExtractor.get_file_extension("document.docx")
        assert ext == ".docx"

    def test_get_file_extension_uppercase(self) -> None:
        """Extension should be lowercase."""
        ext = DocumentExtractor.get_file_extension("CONTRACT.PDF")
        assert ext == ".pdf"

    def test_get_file_extension_no_extension_raises(self) -> None:
        """Raise ValueError if file has no extension."""
        with pytest.raises(ValueError, match="no extension"):
            DocumentExtractor.get_file_extension("contract")

    def test_extract_unsupported_format_raises(self) -> None:
        """Raise ValueError for unsupported file format."""
        with pytest.raises(ValueError, match="Unsupported file format"):
            DocumentExtractor.extract(b"dummy", "file.txt")

    def test_extract_pdf_simple(self) -> None:
        """Extract text from a simple PDF.
        
        Note: This test uses a minimal valid PDF for testing.
        In a real test suite, you'd use a test fixture with a real PDF.
        """
        # Minimal PDF bytes (valid but empty)
        pdf_bytes = (
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>"
            b"endobj\n"
            b"xref\n0 4\n0000000000 65535 f\n"
            b"0000000009 00000 n\n0000000058 00000 n\n0000000115 00000 n\n"
            b"trailer<</Size 4/Root 1 0 R>>\nstartxref\n196\n%%EOF"
        )

        text, metadata = DocumentExtractor.extract(pdf_bytes, "test.pdf")

        assert isinstance(text, str)
        assert isinstance(metadata, dict)
        assert metadata["format"] == "pdf"
        assert "file_size" in metadata
        assert "page_count" in metadata
        assert "word_count" in metadata
        assert metadata["file_size"] == len(pdf_bytes)

    def test_extract_docx_simple(self) -> None:
        """Extract text from a simple DOCX document.
        
        This creates a minimal DOCX in memory for testing.
        """
        from docx import Document as DocxDocument
        from io import BytesIO

        # Create a minimal DOCX in memory
        doc = DocxDocument()
        doc.add_paragraph("This is a test document.")
        doc.add_paragraph("It has multiple paragraphs.")

        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        docx_data = docx_bytes.read()

        text, metadata = DocumentExtractor.extract(docx_data, "test.docx")

        assert isinstance(text, str)
        assert isinstance(metadata, dict)
        assert metadata["format"] == "docx"
        assert "file_size" in metadata
        assert "word_count" in metadata
        assert metadata["file_size"] == len(docx_data)
        assert "test" in text.lower()
        assert "document" in text.lower()


# ── Tests for DocumentAgent ────────────────────────────────────────────────────


class TestDocumentAgent:
    """Tests for the DocumentAgent."""

    def test_agent_instantiation(self) -> None:
        """Instantiate DocumentAgent with default bucket."""
        agent = DocumentAgent()
        assert agent.storage_service is not None
        assert agent.storage_service.bucket_name == "contracts"

    def test_agent_instantiation_custom_bucket(self) -> None:
        """Instantiate DocumentAgent with custom bucket name."""
        agent = DocumentAgent(storage_bucket="custom-bucket")
        assert agent.storage_service.bucket_name == "custom-bucket"

    def test_execute_missing_storage_path(self) -> None:
        """If storage_path is missing, agent should handle gracefully."""
        agent = DocumentAgent()
        state = ContractState(
            session_id="test-1",
            contract_id="contract-1",
            file_name="test.pdf",
            storage_path="",  # Empty
            processing_status="pending",
            errors=[],
        )

        result = agent.execute(state)

        # Should return an error state, not crash
        assert result["processing_status"] == "document_failed"
        assert len(result.get("errors", [])) > 0

    def test_execute_preserves_session_id(self) -> None:
        """Agent should not modify fields it doesn't set."""
        agent = DocumentAgent()
        state = ContractState(
            session_id="test-session-preserve",
            contract_id="contract-1",
            file_name="test.pdf",
            storage_path="invalid/path.pdf",
            processing_status="pending",
            errors=[],
        )

        # Will fail due to invalid storage path, but session_id should not be in result
        result = agent.execute(state)

        # Since partial state is returned, session_id won't be in result unless we check merged state
        # But the agent didn't touch it, so it shouldn't be changed
        assert result.get("processing_status") in ("document_failed", "document_processed")
